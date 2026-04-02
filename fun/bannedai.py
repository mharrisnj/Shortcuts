#!/usr/bin/env python3
"""
Generate a polished .docx report for AI NSFW content creation service blocking/review.

Requirements:
    pip install python-docx

Example:
    python build_ai_nsfw_report.py --client "Acme Corp" --prepared-by "Security Operations"

Notes:
- The document includes a Word table of contents field. If it does not populate
  automatically when opened, press Ctrl+A and then F9 in Microsoft Word.
"""

from __future__ import annotations

import argparse
import sys
from collections import Counter
from datetime import date, datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print(
        "Missing dependency: python-docx\n"
        "Install it with:\n"
        "    pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


SITE_RECORDS: list[dict[str, str]] = [
    {
        "domain": "pornpen.ai",
        "action": "Block",
        "category": "Adult AI image generation",
        "confidence": "High",
        "notes": "Publicly associated with adult-focused AI image generation.",
    },
    {
        "domain": "promptchan.ai",
        "action": "Block",
        "category": "Adult AI image generation",
        "confidence": "High",
        "notes": "Publicly associated with NSFW AI image generation.",
    },
    {
        "domain": "seduced.ai",
        "action": "Block",
        "category": "Adult AI image generation",
        "confidence": "High",
        "notes": "Adult-focused AI image generation positioning.",
    },
    {
        "domain": "soulgen.net",
        "action": "Block",
        "category": "Adult AI image generation",
        "confidence": "High",
        "notes": "Commonly associated with adult character and image generation.",
    },
    {
        "domain": "crushon.ai",
        "action": "Block",
        "category": "Adult AI chat/roleplay",
        "confidence": "High",
        "notes": "Known for uncensored or adult AI chat use cases.",
    },
    {
        "domain": "spicychat.ai",
        "action": "Block",
        "category": "Adult AI chat/roleplay",
        "confidence": "High",
        "notes": "Adult-oriented AI roleplay and chat platform.",
    },
    {
        "domain": "candy.ai",
        "action": "Block",
        "category": "Adult AI companion",
        "confidence": "High",
        "notes": "Adult or romantic AI companion service.",
    },
    {
        "domain": "dreamgf.ai",
        "action": "Block",
        "category": "Adult AI companion",
        "confidence": "High",
        "notes": "Adult or romantic AI companion service.",
    },
    {
        "domain": "muah.ai",
        "action": "Block",
        "category": "Adult AI companion / media generation",
        "confidence": "High",
        "notes": "Adult-oriented AI companion and media generation features.",
    },
    {
        "domain": "girlfriendgpt.com",
        "action": "Block",
        "category": "Adult AI companion",
        "confidence": "High",
        "notes": "Adult or romantic chatbot positioning.",
    },
    {
        "domain": "pephop.ai",
        "action": "Block",
        "category": "Adult AI chat/roleplay",
        "confidence": "High",
        "notes": "Commonly associated with uncensored or adult AI chat.",
    },
    {
        "domain": "civitai.com",
        "action": "Review / Conditional Block",
        "category": "Model hosting",
        "confidence": "Medium",
        "notes": "Mixed-use platform with notable NSFW model and image availability.",
    },
    {
        "domain": "tensor.art",
        "action": "Review / Conditional Block",
        "category": "Model hosting / generation",
        "confidence": "Medium",
        "notes": "Mixed-use platform with mature-content-capable models.",
    },
    {
        "domain": "seaart.ai",
        "action": "Review / Conditional Block",
        "category": "Image generation",
        "confidence": "Medium",
        "notes": "Mixed-use image platform with mature-content availability.",
    },
    {
        "domain": "novelai.net",
        "action": "Review / Conditional Block",
        "category": "Text/image generation",
        "confidence": "Medium",
        "notes": "Mixed-use platform often associated with erotic or anime-style generation.",
    },
    {
        "domain": "chub.ai",
        "action": "Review / Conditional Block",
        "category": "Character hosting",
        "confidence": "Medium",
        "notes": "Mixed-use character hosting platform with adult roleplay availability.",
    },
    {
        "domain": "venus.chub.ai",
        "action": "Review / Conditional Block",
        "category": "Character hosting subdomain",
        "confidence": "Medium",
        "notes": "Notable subdomain associated with adult-oriented use cases.",
    },
    {
        "domain": "janitorai.com",
        "action": "Review / Conditional Block",
        "category": "Chat platform",
        "confidence": "Medium",
        "notes": "Mixed-use chat platform often associated with uncensored or adult bots.",
    },
]

MATCH_PATTERNS: list[str] = [
    "pornpen.ai",
    "*.pornpen.ai",
    "promptchan.ai",
    "*.promptchan.ai",
    "seduced.ai",
    "*.seduced.ai",
    "soulgen.net",
    "*.soulgen.net",
    "crushon.ai",
    "*.crushon.ai",
    "spicychat.ai",
    "*.spicychat.ai",
    "candy.ai",
    "*.candy.ai",
    "dreamgf.ai",
    "*.dreamgf.ai",
    "muah.ai",
    "*.muah.ai",
    "girlfriendgpt.com",
    "*.girlfriendgpt.com",
    "pephop.ai",
    "*.pephop.ai",
    "civitai.com",
    "*.civitai.com",
    "tensor.art",
    "*.tensor.art",
    "seaart.ai",
    "*.seaart.ai",
    "novelai.net",
    "*.novelai.net",
    "chub.ai",
    "*.chub.ai",
    "venus.chub.ai",
    "janitorai.com",
    "*.janitorai.com",
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a .docx report summarizing AI NSFW content creation services."
    )
    parser.add_argument(
        "--client",
        default="[Client Name]",
        help="Client name to display in the report.",
    )
    parser.add_argument(
        "--prepared-by",
        default="[Your Name / Team]",
        help="Author or team name to display in the report.",
    )
    parser.add_argument(
        "--output",
        default="AI_NSFW_Banned_List_Update_Summary_Report.docx",
        help="Output .docx file path.",
    )
    parser.add_argument(
        "--report-date",
        default=date.today().strftime("%d %B %Y"),
        help='Report date, e.g. "28 March 2026".',
    )
    return parser.parse_args()


def set_document_defaults(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)

    # Footer with page number
    for section in document.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = "Page "
        add_field_code(p, "PAGE")


def set_core_properties(
    document: Document,
    title: str,
    author: str,
    subject: str,
) -> None:
    props = document.core_properties
    props.title = title
    props.author = author
    props.subject = subject
    props.comments = "Generated by Python using python-docx."
    props.created = datetime.now()
    props.modified = datetime.now()


def add_field_code(paragraph: Any, field_code: str, placeholder_text: str = " ") -> None:
    """
    Insert a Word field code, e.g. PAGE or TOC.
    """
    run = paragraph.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = field_code

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = placeholder_text

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r.append(fld_begin)
    r.append(instr)
    r.append(fld_separate)
    r.append(text)
    r.append(fld_end)


def enable_update_fields_on_open(document: Document) -> None:
    """
    Ask Word to update fields such as TOC when the document opens.
    """
    settings = document.settings._element
    update_tag = qn("w:updateFields")
    if not any(child.tag == update_tag for child in settings):
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def format_table_header(row: Any, fill: str = "D9EAF7") -> None:
    for cell in row.cells:
        shade_cell(cell, fill)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)


def add_cover_page(
    document: Document,
    report_title: str,
    client: str,
    prepared_by: str,
    report_date: str,
) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n" * 4)

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(report_title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle_p = document.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("Banned List Update Summary Report")
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(14)

    document.add_paragraph()

    meta_lines = [
        f"Prepared for: {client}",
        f"Prepared by: {prepared_by}",
        f"Date: {report_date}",
        "Document Type: Summary Report",
    ]
    for line in meta_lines:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Calibri"
        r.font.size = Pt(12)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n" * 8)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run(
        "This document is intended as a practical policy support artifact and a non-exhaustive seed list."
    )
    r.italic = True
    r.font.size = Pt(10)

    document.add_page_break()


def add_contents_page(document: Document) -> None:
    document.add_heading("Contents", level=1)

    p = document.add_paragraph()
    add_field_code(p, r'TOC \o "1-3" \h \z \u', "Right-click and update the Table of Contents if needed.")

    note = document.add_paragraph()
    note.style = document.styles["Normal"]
    run = note.add_run(
        "If the table of contents does not populate automatically in Word, press Ctrl+A and then F9."
    )
    run.italic = True
    run.font.size = Pt(10)

    document.add_page_break()


def add_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph(text)
    p.style = document.styles["Normal"]


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_summary_metrics(document: Document) -> None:
    total_sites = len(SITE_RECORDS)
    action_counts = Counter(record["action"] for record in SITE_RECORDS)
    confidence_counts = Counter(record["confidence"] for record in SITE_RECORDS)

    document.add_heading("Summary Metrics", level=2)

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    format_table_header(table.rows[0])

    metrics = [
        ("Total candidate domain entries", str(total_sites)),
        ("Immediate block recommendations", str(action_counts.get("Block", 0))),
        (
            "Review / conditional block recommendations",
            str(action_counts.get("Review / Conditional Block", 0)),
        ),
        ("High-confidence entries", str(confidence_counts.get("High", 0))),
        ("Medium-confidence entries", str(confidence_counts.get("Medium", 0))),
        ("Total recommended match patterns", str(len(MATCH_PATTERNS))),
    ]

    for metric, value in metrics:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value


def add_report_body(document: Document) -> None:
    document.add_heading("1. Executive Summary", level=1)
    add_paragraph(
        document,
        "This report summarizes a set of online services that are publicly associated with "
        "AI-enabled NSFW or adult-content creation, including image generation, uncensored "
        "chat and roleplay, and AI companion services. The purpose of this review is to "
        "support an update to the client’s banned or restricted services list.",
    )
    add_paragraph(
        document,
        "The domains identified in this report are grouped into two categories: "
        "(1) high-precision adult-focused services, and "
        "(2) mixed-use platforms with notable NSFW availability. "
        "This two-tier structure helps balance blocking effectiveness against false-positive risk.",
    )
    add_paragraph(
        document,
        "The recommended operating model is to block high-precision adult-focused services "
        "immediately and place mixed-use services into a review, conditional block, or restricted-access category "
        "depending on the client’s policy scope and tolerance for false positives.",
    )

    add_summary_metrics(document)

    document.add_heading("2. Objective", level=1)
    add_paragraph(
        document,
        "The objective of this report is to provide a practical, defensible summary of AI-related "
        "web services that may warrant inclusion in a client banned list where the policy goal is to "
        "reduce or prevent access to AI tools that enable NSFW or adult content creation.",
    )

    document.add_heading("3. Scope", level=1)
    add_paragraph(
        document,
        "This report focuses on web services associated with one or more of the following risk areas:",
    )
    add_bullets(
        document,
        [
            "AI-generated adult or NSFW images",
            "AI-generated erotic or adult-oriented chat and roleplay",
            "AI companion platforms with explicit or adult-oriented functionality",
            "Model-hosting or character-hosting platforms with meaningful NSFW availability",
        ],
    )
    add_paragraph(
        document,
        "This report should be treated as a seed list for compliance, filtering, or policy review rather than "
        "a permanent or exhaustive registry.",
    )

    document.add_heading("4. Methodology", level=1)
    add_paragraph(
        document,
        "Candidate domains were assessed using practical screening criteria such as public branding or "
        "positioning as adult-oriented or uncensored AI, common public association with NSFW AI generation or roleplay, "
        "platform type, and the likelihood that the service would fall within the client’s policy definition of "
        "prohibited AI-enabled NSFW content creation.",
    )
    add_paragraph(
        document,
        "Because these services evolve rapidly, all entries should be treated as subject to periodic verification.",
    )

    document.add_heading("5. Summary of Findings", level=1)

    document.add_heading("5.1 High-Precision Adult-Focused Services", level=2)
    add_paragraph(
        document,
        "These domains are more likely to fit a direct blocking policy with relatively lower false-positive risk. "
        "The main categories identified were adult AI image generation, adult AI chat and roleplay, "
        "and adult AI companion services.",
    )

    document.add_heading("5.2 Mixed-Use Platforms with NSFW Availability", level=2)
    add_paragraph(
        document,
        "These domains may require more nuanced treatment. Although they are not always exclusively adult-focused, "
        "they can still present policy or reputational risk where NSFW AI generation, hosting, or distribution is a concern.",
    )
    add_bullets(
        document,
        [
            "Hosting of NSFW-capable checkpoints, LoRAs, and models",
            "User-generated character libraries containing adult roleplay content",
            "Mature-content generation features coexisting with general-use functionality",
        ],
    )

    document.add_heading("6. Recommended Policy Approach", level=1)

    document.add_heading("6.1 Immediate Block Recommendations", level=2)
    add_bullets(
        document,
        [
            "Adult-focused AI image generators",
            "Explicitly uncensored AI roleplay and chat services",
            "Adult-branded AI companion services",
        ],
    )

    document.add_heading("6.2 Review-Based or Conditional Block Recommendations", level=2)
    add_bullets(
        document,
        [
            "Full block where policy scope is broad and false-positive tolerance is high",
            "Restricted access pending business-owner or manager approval",
            "Monitoring or review queue placement for mixed-use services",
            "Blocking only specific subdomains where technically possible",
        ],
    )

    document.add_heading("6.3 Technical Matching Guidance", level=2)
    add_paragraph(
        document,
        "For enforcement purposes, use both the root domain and wildcard subdomain patterns where supported. "
        "This reduces the chance of missing alternate service endpoints, hosted subdomains, or future product separation.",
    )
    add_bullets(
        document,
        [
            "example.com",
            "*.example.com",
        ],
    )

    document.add_heading("7. Limitations", level=1)
    add_bullets(
        document,
        [
            "Domain ownership and use cases can change over time",
            "Platform moderation practices may change",
            "Some services are mixed-use and may require contextual review",
            "New domains, rebrands, and subdomains may appear without notice",
        ],
    )
    add_paragraph(
        document,
        "For these reasons, a recurring review cycle, such as quarterly reassessment, is recommended.",
    )

    document.add_heading("8. Conclusion", level=1)
    add_paragraph(
        document,
        "The reviewed domains provide a reasonable starting point for updating a banned list focused on "
        "AI-enabled NSFW content creation. A two-tier model is recommended: Tier 1 for immediate blocking "
        "of high-confidence adult-focused domains, and Tier 2 for review or conditional handling of mixed-use "
        "platforms with notable NSFW availability.",
    )
    add_paragraph(
        document,
        "This approach balances policy effectiveness with operational practicality and helps reduce both "
        "underblocking and unnecessary false positives.",
    )


def add_appendix_a(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Appendix A — Candidate Domains for Banned / Review List", level=1)

    add_paragraph(
        document,
        "The following table lists the candidate domains identified during this review, along with the proposed "
        "action, category, confidence level, and a short rationale.",
    )

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Domain", "Suggested Action", "Category", "Confidence", "Notes"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    format_table_header(table.rows[0])

    for record in SITE_RECORDS:
        row = table.add_row().cells
        row[0].text = record["domain"]
        row[1].text = record["action"]
        row[2].text = record["category"]
        row[3].text = record["confidence"]
        row[4].text = record["notes"]


def add_appendix_b(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Appendix B — Recommended Match Patterns", level=1)

    add_paragraph(
        document,
        "Use both root-domain and wildcard subdomain matching where your filtering platform supports it.",
    )

    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "Match Pattern"
    format_table_header(table.rows[0])

    for pattern in MATCH_PATTERNS:
        row = table.add_row().cells
        row[0].text = pattern


def build_report(client: str, prepared_by: str, report_date: str, output_path: Path) -> Path:
    document = Document()
    report_title = "AI NSFW Content Creation Services"

    set_document_defaults(document)
    set_core_properties(
        document=document,
        title=f"{report_title} - Banned List Update Summary Report",
        author=prepared_by,
        subject="Banned List Update Summary Report",
    )
    enable_update_fields_on_open(document)

    add_cover_page(
        document=document,
        report_title=report_title,
        client=client,
        prepared_by=prepared_by,
        report_date=report_date,
    )
    add_contents_page(document)
    add_report_body(document)
    add_appendix_a(document)
    add_appendix_b(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def main() -> None:
    args = parse_args()
    output_path = Path(args.output).expanduser().resolve()

    try:
        saved_to = build_report(
            client=args.client,
            prepared_by=args.prepared_by,
            report_date=args.report_date,
            output_path=output_path,
        )
    except Exception as exc:
        print(f"Failed to generate report: {exc}", file=sys.stderr)
        sys.exit(1)

    print(f"Report created successfully: {saved_to}")


if __name__ == "__main__":
    main()